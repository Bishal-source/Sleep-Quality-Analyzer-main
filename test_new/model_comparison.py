import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report
from xgboost import XGBClassifier
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
data_path = "C:/Users/bishal/Desktop/Sleep-Quality-Analyzer-main/dataset/sleep_dataset_balanced_influence.csv"
output_dir = "C:/Users/bishal/Desktop/Sleep-Quality-Analyzer-main/test_new"

# Load dataset
df = pd.read_csv(data_path)
print("Dataset shape:", df.shape)

# Clean data
df_clean = df.dropna()
print(f"Clean dataset: {df_clean.shape}")

# Features and target
feature_cols = ['Sleep Duration', 'Stress Level', 'Physical Activity', 'Caffeine Intake',
                'Screen Time', 'SpO2 Before', 'SpO2 After', 'Heart Rate Before', 'Heart Rate After']
target_col = 'Sleep Quality'

X = df_clean[feature_cols]
y = df_clean[target_col]

# Convert to binary classification (Good: >=5, Poor: <5)
y_binary = (y >= 5).astype(int)
print(f"Target distribution: {np.bincount(y_binary)}")

# Split data
X_train, X_test, y_train, y_test = train_test_split(X, y_binary, test_size=0.2, random_state=42, stratify=y_binary)
print(f"Train: {X_train.shape}, Test: {X_test.shape}")

# =====================================================
# MODEL 1: WITHOUT PCA
# =====================================================
print("\n" + "="*60)
print("TRAINING XGBOOST WITHOUT PCA")
print("="*60)

scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

model_no_pca = XGBClassifier(
    n_estimators=100,
    max_depth=6,
    learning_rate=0.1,
    random_state=42,
    use_label_encoder=False,
    eval_metric='logloss'
)

model_no_pca.fit(X_train_scaled, y_train)
y_pred_no_pca = model_no_pca.predict(X_test_scaled)

# Metrics without PCA
metrics_no_pca = {
    'Accuracy': accuracy_score(y_test, y_pred_no_pca),
    'Precision': precision_score(y_test, y_pred_no_pca),
    'Recall': recall_score(y_test, y_pred_no_pca),
    'F1 Score': f1_score(y_test, y_pred_no_pca)
}

print("\nWithout PCA:")
for metric, value in metrics_no_pca.items():
    print(f"  {metric}: {value:.4f}")

print("\nClassification Report:")
print(classification_report(y_test, y_pred_no_pca))

# Confusion matrix
cm_no_pca = confusion_matrix(y_test, y_pred_no_pca)
print("Confusion Matrix:")
print(cm_no_pca)

# =====================================================
# MODEL 2: WITH PCA (using top 5 components - 77% variance)
# =====================================================
print("\n" + "="*60)
print("TRAINING XGBOOST WITH PCA")
print("="*60)

# Scale first
X_train_scaled_all = scaler.fit_transform(X_train)
X_test_scaled_all = scaler.transform(X_test)

# PCA with 5 components (captures ~77% variance)
pca = PCA(n_components=5)
X_train_pca = pca.fit_transform(X_train_scaled_all)
X_test_pca = pca.transform(X_test_scaled_all)

print(f"PCA components: {pca.n_components_}")
print(f"Explained variance: {sum(pca.explained_variance_ratio_)*100:.2f}%")

model_with_pca = XGBClassifier(
    n_estimators=100,
    max_depth=6,
    learning_rate=0.1,
    random_state=42,
    use_label_encoder=False,
    eval_metric='logloss'
)

model_with_pca.fit(X_train_pca, y_train)
y_pred_pca = model_with_pca.predict(X_test_pca)

# Metrics with PCA
metrics_pca = {
    'Accuracy': accuracy_score(y_test, y_pred_pca),
    'Precision': precision_score(y_test, y_pred_pca),
    'Recall': recall_score(y_test, y_pred_pca),
    'F1 Score': f1_score(y_test, y_pred_pca)
}

print("\nWith PCA (5 components):")
for metric, value in metrics_pca.items():
    print(f"  {metric}: {value:.4f}")

print("\nClassification Report:")
print(classification_report(y_test, y_pred_pca))

# Confusion matrix
cm_pca = confusion_matrix(y_test, y_pred_pca)
print("Confusion Matrix:")
print(cm_pca)

# =====================================================
# CREATE VISUALIZATIONS
# =====================================================
fig, axes = plt.subplots(1, 2, figsize=(14, 5))

# Plot 1: Metrics comparison
metrics_names = list(metrics_no_pca.keys())
x = np.arange(len(metrics_names))
width = 0.35

bars1 = axes[0].bar(x - width/2, list(metrics_no_pca.values()), width, label='Without PCA', color='steelblue')
bars2 = axes[0].bar(x + width/2, list(metrics_pca.values()), width, label='With PCA', color='coral')

axes[0].set_ylabel('Score')
axes[0].set_title('Model Performance Comparison')
axes[0].set_xticks(x)
axes[0].set_xticklabels(metrics_names)
axes[0].legend()
axes[0].set_ylim(0, 1)

for bar in bars1:
    axes[0].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02, f'{bar.get_height():.3f}', ha='center', va='bottom', fontsize=9)
for bar in bars2:
    axes[0].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02, f'{bar.get_height():.3f}', ha='center', va='bottom', fontsize=9)

# Plot 2: Explained variance
axes[1].bar(range(1, 6), pca.explained_variance_ratio_, alpha=0.7, color='green', label='Individual')
axes[1].plot(range(1, 6), np.cumsum(pca.explained_variance_ratio_), 'ro-', label='Cumulative')
axes[1].axhline(y=0.95, color='g', linestyle='--', alpha=0.5)
axes[1].set_xlabel('Principal Component')
axes[1].set_ylabel('Explained Variance')
axes[1].set_title('PCA Explained Variance (Top 5 Components)')
axes[1].legend()
axes[1].set_xticks(range(1, 6))

plt.tight_layout()
plt.savefig(os.path.join(output_dir, "model_comparison.png"), dpi=150)
print(f"\nVisualization saved to: {os.path.join(output_dir, 'model_comparison.png')}")

# =====================================================
# CREATE WORD DOCUMENT REPORT
# =====================================================
doc = Document()

# Title
title = doc.add_heading('Sleep Quality Analysis: XGBoost Model Comparison Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This report compares the performance of XGBoost classifiers trained on sleep quality data '
    'with and without Principal Component Analysis (PCA) dimensionality reduction. The dataset '
    'contains various physiological and lifestyle factors that influence sleep quality.'
)

# Dataset Description
doc.add_heading('2. Dataset Description', level=1)
doc.add_paragraph(f'Original dataset shape: {df.shape[0]} rows, {df.shape[1]} columns')
doc.add_paragraph(f'After removing missing values: {df_clean.shape[0]} rows')
doc.add_paragraph(f'Features used: {len(feature_cols)}')
doc.add_paragraph(f'Features: {", ".join(feature_cols)}')

doc.add_paragraph('\nTarget variable: Sleep Quality (converted to binary: Good >= 5, Poor < 5)')
doc.add_paragraph(f'Class distribution: Good: {sum(y_binary)}, Poor: {len(y_binary) - sum(y_binary)}')

doc.add_paragraph(f'\nTrain/Test split: 80%/20% (Train: {X_train.shape[0]}, Test: {X_test.shape[0]})')

# PCA Description
doc.add_heading('3. PCA Analysis', level=1)
doc.add_paragraph(f'PCA was applied with {pca.n_components_} components.')
doc.add_paragraph(f'Total explained variance: {sum(pca.explained_variance_ratio_)*100:.2f}%')
doc.add_paragraph('Individual component variances:')
for i, var in enumerate(pca.explained_variance_ratio_):
    doc.add_paragraph(f'  PC{i+1}: {var*100:.2f}%', style='List Bullet')

# Results
doc.add_heading('4. Model Performance Results', level=1)

# Table: Metrics Comparison
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Without PCA'
header_cells[2].text = 'With PCA'

metrics_list = ['Accuracy', 'Precision', 'Recall', 'F1 Score']
for i, metric in enumerate(metrics_list):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = metric
    row_cells[1].text = f'{metrics_no_pca[metric]:.4f}'
    row_cells[2].text = f'{metrics_pca[metric]:.4f}'

doc.add_paragraph()

# Without PCA Details
doc.add_heading('4.1 Without PCA Model', level=2)
doc.add_paragraph('This model uses all 9 original features without any dimensionality reduction.')
for metric, value in metrics_no_pca.items():
    doc.add_paragraph(f'{metric}: {value:.4f}', style='List Bullet')

doc.add_paragraph('\nConfusion Matrix (Without PCA):')
cm_text_no_pca = f'TN={cm_no_pca[0,0]}, FP={cm_no_pca[0,1]}, FN={cm_no_pca[1,0]}, TP={cm_no_pca[1,1]}'
doc.add_paragraph(cm_text_no_pca)

# With PCA Details
doc.add_heading('4.2 With PCA Model', level=2)
doc.add_paragraph(f'This model uses {pca.n_components_} principal components capturing {sum(pca.explained_variance_ratio_)*100:.2f}% of variance.')
for metric, value in metrics_pca.items():
    doc.add_paragraph(f'{metric}: {value:.4f}', style='List Bullet')

doc.add_paragraph('\nConfusion Matrix (With PCA):')
cm_text_pca = f'TN={cm_pca[0,0]}, FP={cm_pca[0,1]}, FN={cm_pca[1,0]}, TP={cm_pca[1,1]}'
doc.add_paragraph(cm_text_pca)

# Analysis
doc.add_heading('5. Analysis & Conclusion', level=1)

accuracy_diff = metrics_no_pca['Accuracy'] - metrics_pca['Accuracy']
doc.add_paragraph(f'Difference in Accuracy (Without PCA - With PCA): {accuracy_diff:.4f}')

if metrics_no_pca['Accuracy'] > metrics_pca['Accuracy']:
    doc.add_paragraph(
        f'The model WITHOUT PCA performs slightly better with {metrics_no_pca["Accuracy"]*100:.2f}% accuracy. '
        'This suggests that all original features contain valuable information for predicting sleep quality.'
    )
else:
    doc.add_paragraph(
        f'The model WITH PCA performs slightly better with {metrics_pca["Accuracy"]*100:.2f}% accuracy. '
        'This suggests that dimensionality reduction helped remove noise and improved generalization.'
    )

doc.add_paragraph(
    '\nPCA reduced the feature space from 9 dimensions to 5 dimensions while retaining '
    f'{sum(pca.explained_variance_ratio_)*100:.2f}% of the variance. This can be beneficial for:')
doc.add_paragraph('  - Faster model training', style='List Bullet')
doc.add_paragraph('  - Reduced overfitting', style='List Bullet')
doc.add_paragraph('  - Better generalization to new data', style='List Bullet')

# Add visualization
doc.add_heading('6. Visualizations', level=1)
doc.add_picture(os.path.join(output_dir, "model_comparison.png"), width=Inches(6))

# Save document
doc.save(os.path.join(output_dir, "model_comparison_report.docx"))
print(f"\nWord document saved to: {os.path.join(output_dir, 'model_comparison_report.docx')}")

print("\n" + "="*60)
print("ANALYSIS COMPLETE!")
print("="*60)