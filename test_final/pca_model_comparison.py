import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.decomposition import PCA
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report
from xgboost import XGBClassifier
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
data_path = "C:/Users/bishal/Desktop/Sleep-Quality-Analyzer-main/dataset/sleep_quality_dataset.csv"
output_dir = "C:/Users/bishal/Desktop/Sleep-Quality-Analyzer-main/test_final"

print("="*70)
print("SLEEP QUALITY DATASET - PCA ANALYSIS & MODEL COMPARISON")
print("="*70)

# Load dataset
df = pd.read_csv(data_path)
print(f"\nDataset shape: {df.shape}")
print(f"Columns: {df.columns.tolist()}")
print(f"\nFirst few rows:")
print(df.head())

# Check target distribution
print(f"\nTarget distribution (Sleep Quality):")
print(df['Sleep Quality'].value_counts())

# Check missing values
print(f"\nMissing values:")
print(df.isnull().sum())

# Clean data - drop rows with missing values
df_clean = df.dropna()
print(f"\nAfter cleaning: {df_clean.shape[0]} rows")

# Features and target
feature_cols = ['Sleep Duration', 'Stress Level', 'Physical Activity', 'Caffeine Intake',
                'Screen Time', 'SpO2 Before', 'SpO2 After', 'Heart Rate Before', 'Heart Rate After']
target_col = 'Sleep Quality'

X = df_clean[feature_cols]
y = df_clean[target_col]

# Encode target
le = LabelEncoder()
y_encoded = le.fit_transform(y)
print(f"\nTarget classes: {le.classes_}")
print(f"Encoded: {dict(zip(le.classes_, range(len(le.classes_))))}")

# Split data
X_train, X_test, y_train, y_test = train_test_split(
    X, y_encoded, test_size=0.2, random_state=42, stratify=y_encoded
)
print(f"\nTrain: {X_train.shape}, Test: {X_test.shape}")
print(f"Train target distribution: {np.bincount(y_train)}")
print(f"Test target distribution: {np.bincount(y_test)}")

# =====================================================
# PCA ANALYSIS
# =====================================================
print("\n" + "="*70)
print("PCA ANALYSIS")
print("="*70)

# Standardize features
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# Perform PCA with all components first to see variance
pca_full = PCA()
pca_full.fit(X_train_scaled)

print("\nExplained Variance Ratio (all components):")
for i, (var, cum) in enumerate(zip(pca_full.explained_variance_ratio_, np.cumsum(pca_full.explained_variance_ratio_))):
    print(f"  PC{i+1}: {var*100:.2f}% (Cumulative: {cum*100:.2f}%)")

# Determine optimal number of components (>= 80% variance)
n_components = np.argmax(np.cumsum(pca_full.explained_variance_ratio_) >= 0.80) + 1
print(f"\nOptimal components (>=80% variance): {n_components}")

# Apply PCA with optimal components
pca = PCA(n_components=n_components)
X_train_pca = pca.fit_transform(X_train_scaled)
X_test_pca = pca.transform(X_test_scaled)

print(f"PCA applied: {pca.n_components_} components")
print(f"Explained variance: {sum(pca.explained_variance_ratio_)*100:.2f}%")

# Save PCA results
pca_results_df = pd.DataFrame(X_train_pca, columns=[f'PC{i+1}' for i in range(n_components)])
pca_results_df['Sleep Quality'] = le.inverse_transform(y_train)
pca_results_df.to_csv(os.path.join(output_dir, "pca_transformed_data.csv"), index=False)
print(f"PCA results saved to: {os.path.join(output_dir, 'pca_transformed_data.csv')}")

# Save explained variance
variance_df = pd.DataFrame({
    'Component': [f'PC{i+1}' for i in range(len(pca_full.explained_variance_ratio_))],
    'Explained_Variance': pca_full.explained_variance_ratio_,
    'Cumulative_Variance': np.cumsum(pca_full.explained_variance_ratio_)
})
variance_df.to_csv(os.path.join(output_dir, "explained_variance.csv"), index=False)

# Save loadings
loadings = pd.DataFrame(
    pca_full.components_.T,
    columns=[f'PC{i+1}' for i in range(len(pca_full.explained_variance_ratio_))],
    index=feature_cols
)
loadings.to_csv(os.path.join(output_dir, "pca_loadings.csv"))

# =====================================================
# MODEL 1: XGBOOST WITHOUT PCA
# =====================================================
print("\n" + "="*70)
print("TRAINING XGBOOST WITHOUT PCA")
print("="*70)

model_no_pca = XGBClassifier(
    n_estimators=100,
    max_depth=6,
    learning_rate=0.1,
    random_state=42,
    eval_metric='mlogloss',
    verbosity=0
)

model_no_pca.fit(X_train_scaled, y_train)
y_pred_no_pca = model_no_pca.predict(X_test_scaled)

# Metrics without PCA
metrics_no_pca = {
    'Accuracy': accuracy_score(y_test, y_pred_no_pca),
    'Precision (Macro)': precision_score(y_test, y_pred_no_pca, average='macro'),
    'Recall (Macro)': recall_score(y_test, y_pred_no_pca, average='macro'),
    'F1 Score (Macro)': f1_score(y_test, y_pred_no_pca, average='macro'),
    'Precision (Weighted)': precision_score(y_test, y_pred_no_pca, average='weighted'),
    'Recall (Weighted)': recall_score(y_test, y_pred_no_pca, average='weighted'),
    'F1 Score (Weighted)': f1_score(y_test, y_pred_no_pca, average='weighted')
}

print("\nWithout PCA Metrics:")
for metric, value in metrics_no_pca.items():
    print(f"  {metric}: {value:.4f}")

print("\nClassification Report (Without PCA):")
print(classification_report(y_test, y_pred_no_pca, target_names=le.classes_))

cm_no_pca = confusion_matrix(y_test, y_pred_no_pca)
print("Confusion Matrix (Without PCA):")
print(cm_no_pca)

# =====================================================
# MODEL 2: XGBOOST WITH PCA
# =====================================================
print("\n" + "="*70)
print("TRAINING XGBOOST WITH PCA")
print("="*70)

model_with_pca = XGBClassifier(
    n_estimators=100,
    max_depth=6,
    learning_rate=0.1,
    random_state=42,
    eval_metric='mlogloss',
    verbosity=0
)

model_with_pca.fit(X_train_pca, y_train)
y_pred_pca = model_with_pca.predict(X_test_pca)

# Metrics with PCA
metrics_pca = {
    'Accuracy': accuracy_score(y_test, y_pred_pca),
    'Precision (Macro)': precision_score(y_test, y_pred_pca, average='macro'),
    'Recall (Macro)': recall_score(y_test, y_pred_pca, average='macro'),
    'F1 Score (Macro)': f1_score(y_test, y_pred_pca, average='macro'),
    'Precision (Weighted)': precision_score(y_test, y_pred_pca, average='weighted'),
    'Recall (Weighted)': recall_score(y_test, y_pred_pca, average='weighted'),
    'F1 Score (Weighted)': f1_score(y_test, y_pred_pca, average='weighted')
}

print("\nWith PCA Metrics:")
for metric, value in metrics_pca.items():
    print(f"  {metric}: {value:.4f}")

print("\nClassification Report (With PCA):")
print(classification_report(y_test, y_pred_pca, target_names=le.classes_))

cm_pca = confusion_matrix(y_test, y_pred_pca)
print("Confusion Matrix (With PCA):")
print(cm_pca)

# =====================================================
# CREATE VISUALIZATIONS
# =====================================================
fig, axes = plt.subplots(2, 2, figsize=(14, 12))

# Plot 1: Metrics comparison (Accuracy, Precision, Recall, F1)
metrics_names = ['Accuracy', 'Precision (Macro)', 'Recall (Macro)', 'F1 Score (Macro)']
x = np.arange(len(metrics_names))
width = 0.35

bars1 = axes[0, 0].bar(x - width/2, [metrics_no_pca[m] for m in metrics_names], width, label='Without PCA', color='steelblue')
bars2 = axes[0, 0].bar(x + width/2, [metrics_pca[m] for m in metrics_names], width, label='With PCA', color='coral')

axes[0, 0].set_ylabel('Score')
axes[0, 0].set_title('Model Performance Comparison')
axes[0, 0].set_xticks(x)
axes[0, 0].set_xticklabels(metrics_names, rotation=15)
axes[0, 0].legend()
axes[0, 0].set_ylim(0, 1)

for bar in bars1:
    axes[0, 0].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, f'{bar.get_height():.3f}', ha='center', va='bottom', fontsize=8)
for bar in bars2:
    axes[0, 0].text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01, f'{bar.get_height():.3f}', ha='center', va='bottom', fontsize=8)

# Plot 2: Scree plot
axes[0, 1].bar(range(1, 10), pca_full.explained_variance_ratio_, alpha=0.7, color='green', label='Individual')
axes[0, 1].plot(range(1, 10), np.cumsum(pca_full.explained_variance_ratio_), 'ro-', label='Cumulative')
axes[0, 1].axhline(y=0.80, color='g', linestyle='--', alpha=0.5, label='80% threshold')
axes[0, 1].set_xlabel('Principal Component')
axes[0, 1].set_ylabel('Explained Variance')
axes[0, 1].set_title('PCA Scree Plot')
axes[0, 1].legend()
axes[0, 1].set_xticks(range(1, 10))

# Plot 3: Confusion Matrix Without PCA
im1 = axes[1, 0].imshow(cm_no_pca, cmap='Blues')
axes[1, 0].set_title('Confusion Matrix (Without PCA)')
axes[1, 0].set_xticks(range(len(le.classes_)))
axes[1, 0].set_yticks(range(len(le.classes_)))
axes[1, 0].set_xticklabels(le.classes_)
axes[1, 0].set_yticklabels(le.classes_)
axes[1, 0].set_xlabel('Predicted')
axes[1, 0].set_ylabel('Actual')
for i in range(len(le.classes_)):
    for j in range(len(le.classes_)):
        axes[1, 0].text(j, i, cm_no_pca[i, j], ha='center', va='center', color='white' if cm_no_pca[i, j] > cm_no_pca.max()/2 else 'black')
plt.colorbar(im1, ax=axes[1, 0])

# Plot 4: Confusion Matrix With PCA
im2 = axes[1, 1].imshow(cm_pca, cmap='Oranges')
axes[1, 1].set_title('Confusion Matrix (With PCA)')
axes[1, 1].set_xticks(range(len(le.classes_)))
axes[1, 1].set_yticks(range(len(le.classes_)))
axes[1, 1].set_xticklabels(le.classes_)
axes[1, 1].set_yticklabels(le.classes_)
axes[1, 1].set_xlabel('Predicted')
axes[1, 1].set_ylabel('Actual')
for i in range(len(le.classes_)):
    for j in range(len(le.classes_)):
        axes[1, 1].text(j, i, cm_pca[i, j], ha='center', va='center', color='white' if cm_pca[i, j] > cm_pca.max()/2 else 'black')
plt.colorbar(im2, ax=axes[1, 1])

plt.tight_layout()
plt.savefig(os.path.join(output_dir, "analysis_visualization.png"), dpi=150)
print(f"\nVisualization saved to: {os.path.join(output_dir, 'analysis_visualization.png')}")

# =====================================================
# CREATE WORD DOCUMENT REPORT
# =====================================================
doc = Document()

# Title
title = doc.add_heading('Sleep Quality Analysis Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('PCA Analysis and XGBoost Model Comparison')
doc.add_paragraph(f'Generated: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This report presents a comprehensive analysis of the Sleep Quality Dataset, '
    'including Principal Component Analysis (PCA) for dimensionality reduction and '
    'comparison of XGBoost classifier performance with and without PCA.'
)

# 2. Dataset Description
doc.add_heading('2. Dataset Description', level=1)
doc.add_paragraph(f'Original dataset shape: {df.shape[0]} rows, {df.shape[1]} columns')
doc.add_paragraph(f'After removing missing values: {df_clean.shape[0]} rows')

doc.add_paragraph('\nFeatures (9):')
for col in feature_cols:
    doc.add_paragraph(col, style='List Bullet')

doc.add_paragraph(f'\nTarget variable: Sleep Quality (Categorical)')
doc.add_paragraph(f'Classes: {list(le.classes_)}')
doc.add_paragraph('\nClass distribution:')
for cls in le.classes_:
    count = (df_clean[target_col] == cls).sum()
    doc.add_paragraph(f'  {cls}: {count} ({count/len(df_clean)*100:.1f}%)', style='List Bullet')

doc.add_paragraph(f'\nTrain/Test split: 80%/20%')
doc.add_paragraph(f'Train size: {X_train.shape[0]} samples')
doc.add_paragraph(f'Test size: {X_test.shape[0]} samples')

# 3. PCA Analysis
doc.add_heading('3. PCA Analysis', level=1)
doc.add_paragraph(f'Total principal components: {len(pca_full.explained_variance_ratio_)}')
doc.add_paragraph(f'Components selected for model: {pca.n_components_} (explaining {sum(pca.explained_variance_ratio_)*100:.2f}% variance)')

doc.add_paragraph('\nExplained Variance by Component:')
for i, var in enumerate(pca_full.explained_variance_ratio_):
    doc.add_paragraph(f'  PC{i+1}: {var*100:.2f}%', style='List Bullet')

doc.add_paragraph('\nCumulative Variance:')
for i, cum in enumerate(np.cumsum(pca_full.explained_variance_ratio_)):
    if i < 5 or cum >= 0.80:
        doc.add_paragraph(f'  PC1-PC{i+1}: {cum*100:.2f}%', style='List Bullet')

# 4. Model Performance Results
doc.add_heading('4. Model Performance Results', level=1)

# Table: Metrics Comparison
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Metric'
header_cells[1].text = 'Without PCA'
header_cells[2].text = 'With PCA'

metrics_list = ['Accuracy', 'Precision (Macro)', 'Recall (Macro)', 'F1 Score (Macro)',
                'Precision (Weighted)', 'Recall (Weighted)', 'F1 Score (Weighted)']

for i, metric in enumerate(metrics_list):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = metric
    row_cells[1].text = f'{metrics_no_pca[metric]:.4f}'
    row_cells[2].text = f'{metrics_pca[metric]:.4f}'

doc.add_paragraph()

# 5. Detailed Results
doc.add_heading('5. Detailed Results', level=1)

doc.add_heading('5.1 Without PCA Model', level=2)
doc.add_paragraph('This model uses all 9 original features without dimensionality reduction.')
doc.add_paragraph('Confusion Matrix:')
doc.add_paragraph(f'  True Bad: {cm_no_pca[0,0]}, False Average: {cm_no_pca[0,1]}, False Good: {cm_no_pca[0,2]}')
doc.add_paragraph(f'  False Bad: {cm_no_pca[1,0]}, True Average: {cm_no_pca[1,1]}, False Good: {cm_no_pca[1,2]}')
doc.add_paragraph(f'  False Bad: {cm_no_pca[2,0]}, False Average: {cm_no_pca[2,1]}, True Good: {cm_no_pca[2,2]}')

doc.add_heading('5.2 With PCA Model', level=2)
doc.add_paragraph(f'This model uses {pca.n_components_} principal components explaining {sum(pca.explained_variance_ratio_)*100:.2f}% variance.')
doc.add_paragraph('Confusion Matrix:')
doc.add_paragraph(f'  True Bad: {cm_pca[0,0]}, False Average: {cm_pca[0,1]}, False Good: {cm_pca[0,2]}')
doc.add_paragraph(f'  False Bad: {cm_pca[1,0]}, True Average: {cm_pca[1,1]}, False Good: {cm_pca[1,2]}')
doc.add_paragraph(f'  False Bad: {cm_pca[2,0]}, False Average: {cm_pca[2,1]}, True Good: {cm_pca[2,2]}')

# 6. Analysis & Conclusion
doc.add_heading('6. Analysis & Conclusion', level=1)

accuracy_diff = metrics_no_pca['Accuracy'] - metrics_pca['Accuracy']
doc.add_paragraph(f'Accuracy Difference (Without PCA - With PCA): {accuracy_diff:.4f} ({accuracy_diff*100:.2f}%)')

if metrics_no_pca['Accuracy'] > metrics_pca['Accuracy']:
    doc.add_paragraph(
        f'\nThe model WITHOUT PCA performs better with {metrics_no_pca["Accuracy"]*100:.2f}% accuracy compared to {metrics_pca["Accuracy"]*100:.2f}% with PCA.'
    )
    doc.add_paragraph(
        'This indicates that all original features contain valuable information for predicting sleep quality, '
        'and dimensionality reduction results in some information loss.'
    )
else:
    doc.add_paragraph(
        f'\nThe model WITH PCA performs better with {metrics_pca["Accuracy"]*100:.2f}% accuracy compared to {metrics_no_pca["Accuracy"]*100:.2f}% without PCA.'
    )
    doc.add_paragraph(
        'This suggests that PCA helped remove noise and improved model generalization.'
    )

doc.add_paragraph('\nKey Observations:')
doc.add_paragraph(f'  - Original features: {len(feature_cols)}', style='List Bullet')
doc.add_paragraph(f'  - PCA components used: {pca.n_components_}', style='List Bullet')
doc.add_paragraph(f'  - Variance retained: {sum(pca.explained_variance_ratio_)*100:.2f}%', style='List Bullet')
doc.add_paragraph(f'  - Accuracy improvement without PCA: {accuracy_diff*100:.2f}%', style='List Bullet')

# 7. Visualizations
doc.add_heading('7. Visualizations', level=1)
doc.add_picture(os.path.join(output_dir, "analysis_visualization.png"), width=Inches(6))

# Save document
doc.save(os.path.join(output_dir, "sleep_quality_analysis_report.docx"))
print(f"\nWord document saved to: {os.path.join(output_dir, 'sleep_quality_analysis_report.docx')}")

print("\n" + "="*70)
print("ANALYSIS COMPLETE!")
print("="*70)

# Summary
print("\n" + "="*70)
print("SUMMARY")
print("="*70)
print(f"\nDataset: {df_clean.shape[0]} samples, {len(feature_cols)} features")
print(f"Classes: {list(le.classes_)}")
print(f"\nPCA: {pca.n_components_} components, {sum(pca.explained_variance_ratio_)*100:.2f}% variance")
print(f"\nModel Performance:")
print(f"  Without PCA - Accuracy: {metrics_no_pca['Accuracy']:.4f}")
print(f"  With PCA    - Accuracy: {metrics_pca['Accuracy']:.4f}")
print(f"\nFiles created in {output_dir}:")
for f in os.listdir(output_dir):
    print(f"  - {f}")