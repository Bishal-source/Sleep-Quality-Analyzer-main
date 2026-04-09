"""
Create Word Document with Model Comparison Results
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Load results
results_df = pd.read_csv('model_comparison_results.csv')
results_df = results_df.sort_values('Accuracy', ascending=False)

# Create document
doc = Document()

# Title
title = doc.add_heading('Sleep Quality Model Comparison Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'This report compares the performance of machine learning models for Sleep Quality classification. '
    'The dataset contains 30,000 samples with features including Sleep Duration, Stress Level, '
    'Physical Activity, Caffeine Cups, Screen Time, SpO2 levels, Heart Rate, and Age.'
)

doc.add_paragraph(f'Dataset Size: 28,329 samples (after cleaning)')
doc.add_paragraph('Target Classes: poor, good, best')
doc.add_paragraph('Train/Test Split: 80%/20%')

# Model Comparison Table
doc.add_heading('Model Comparison Summary', level=1)

# Create table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Header row
headers = ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'ROC-AUC']
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    hdr_cells[i].paragraphs[0].runs[0].bold = True

# Data rows
for _, row in results_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Model']
    row_cells[1].text = f"{row['Accuracy']*100:.2f}%"
    row_cells[2].text = f"{row['Precision']*100:.2f}%"
    row_cells[3].text = f"{row['Recall']*100:.2f}%"
    row_cells[4].text = f"{row['F1-Score']*100:.2f}%"
    row_cells[5].text = f"{row['ROC-AUC']*100:.2f}%"

# Best Model Section
doc.add_heading('Best Model: XGBoost', level=1)

doc.add_paragraph(
    'XGBoost (eXtreme Gradient Boosting) is the best performing model for this sleep quality classification task.'
)

doc.add_paragraph('Key Highlights:', style='Heading 3')
doc.add_paragraph(f"• Accuracy: 97.88%")
doc.add_paragraph(f"• Precision: 97.87%")
doc.add_paragraph(f"• Recall: 97.88%")
doc.add_paragraph(f"• F1-Score: 97.87%")
doc.add_paragraph(f"• ROC-AUC: 99.82%")
doc.add_paragraph(f"• Cross-Validation: 97.48% (+/- 0.20%)")

doc.add_paragraph('Why XGBoost?', style='Heading 3')
doc.add_paragraph(
    'XGBoost uses gradient boosting framework which builds trees sequentially, where each new tree corrects errors from previous trees. '
    'It handles missing values well, provides built-in regularization to prevent overfitting, and is highly scalable.'
)

# Model Rankings
doc.add_heading('Model Rankings (by Accuracy)', level=1)

rankings = [
    ('1', 'XGBoost', '97.88%', 'Best overall performance with highest accuracy and F1-score'),
    ('2', 'Decision Tree', '93.08%', 'Good interpretability but prone to overfitting'),
    ('3', 'Random Forest', '92.57%', 'Ensemble method with good ROC-AUC'),
    ('4', 'KNN', '89.89%', 'Simple but less accurate for this dataset')
]

for rank, model, acc, note in rankings:
    doc.add_paragraph(f"{rank}. {model} - {acc}")
    doc.add_paragraph(f"   {note}")

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'XGBoost is recommended as the best model for Sleep Quality prediction. '
    'It achieves the highest accuracy (97.88%) with excellent precision and recall across all classes. '
    'The model shows consistent performance in cross-validation, indicating good generalization capability.'
)

# Save document
doc.save('model_comparison_report.docx')
print("Word document saved as: model_comparison_report.docx")