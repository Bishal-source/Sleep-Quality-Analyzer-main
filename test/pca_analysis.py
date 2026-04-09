"""
PCA Analysis: Feature Importance, Correlation, and Dimensionality Reduction
"""

import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.decomposition import PCA
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from xgboost import XGBClassifier

import matplotlib
matplotlib.use('agg')
import matplotlib.pyplot as plt
import seaborn as sns

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("=" * 70)
print("PCA ANALYSIS FOR SLEEP QUALITY DATASET")
print("=" * 70)

# Load dataset
df = pd.read_csv(r"../dataset/sleep_realistic_final.csv")

# Data cleaning
df_clean = df[
    (df['Sleep Duration'] > 0) &
    (df['SpO2 Before'] >= 0) & (df['SpO2 Before'] <= 100) &
    (df['SpO2 After'] >= 0) & (df['SpO2 After'] <= 100) &
    (df['Heart Rate Before'] > 0) & (df['Heart Rate After'] > 0)
].copy()

# Feature engineering
df_clean['HR_Change'] = df_clean['Heart Rate After'] - df_clean['Heart Rate Before']
df_clean['SpO2_Change'] = df_clean['SpO2 After'] - df_clean['SpO2 Before']

# Encode target
quality_mapping = {'poor': 0, 'good': 1, 'best': 2}
df_clean['target'] = df_clean['Sleep Quality'].map(quality_mapping)

print(f"\nDataset shape: {df_clean.shape}")
print(f"Features: {list(df_clean.columns)}")

# ============== 1. CORRELATION ANALYSIS ==============
print("\n" + "=" * 70)
print("1. CORRELATION ANALYSIS")
print("=" * 70)

features = ['Sleep Duration', 'Stress Level', 'Physical Activity',
            'Caffeine Cups', 'Screen Time', 'SpO2 Before', 'SpO2 After',
            'Heart Rate Before', 'Heart Rate After', 'Age', 'HR_Change', 'SpO2_Change']

X = df_clean[features]
y = df_clean['target']

# Correlation matrix
corr_matrix = X.corr()

# Find highly correlated features (|r| > 0.8)
print("\nHighly Correlated Features (|correlation| > 0.8):")
high_corr_pairs = []
for i in range(len(features)):
    for j in range(i+1, len(features)):
        if abs(corr_matrix.iloc[i, j]) > 0.8:
            high_corr_pairs.append((features[i], features[j], corr_matrix.iloc[i, j]))
            print(f"  {features[i]} <-> {features[j]}: {corr_matrix.iloc[i, j]:.4f}")

# Correlation with target
corr_with_target = df_clean[features + ['target']].corr()['target'].drop('target')
corr_with_target = corr_with_target.sort_values(ascending=False)

print("\nCorrelation with Sleep Quality:")
for feat, corr in corr_with_target.items():
    print(f"  {feat}: {corr:.4f}")

# ============== 2. PCA ANALYSIS ==============
print("\n" + "=" * 70)
print("2. PCA ANALYSIS")
print("=" * 70)

# Scale features
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# Full PCA
pca_full = PCA()
pca_full.fit(X_scaled)

# Explained variance
explained_variance = pca_full.explained_variance_ratio_
cumulative_variance = np.cumsum(explained_variance)

print("\nExplained Variance by Component:")
for i, (var, cum_var) in enumerate(zip(explained_variance, cumulative_variance)):
    print(f"  PC{i+1}: {var*100:.2f}% (Cumulative: {cum_var*100:.2f}%)")

# Find number of components for 95% variance
n_components_95 = np.argmax(cumulative_variance >= 0.95) + 1
n_components_90 = np.argmax(cumulative_variance >= 0.90) + 1

print(f"\nComponents needed for 90% variance: {n_components_90}")
print(f"Components needed for 95% variance: {n_components_95}")

# PCA loadings (feature importance per component)
pca_loadings = pd.DataFrame(
    pca_full.components_.T,
    columns=[f'PC{i+1}' for i in range(len(features))],
    index=features
)

print("\nTop Features per Principal Component:")
for i in range(min(3, len(features))):
    pc_name = f'PC{i+1}'
    top_features = pca_loadings[pc_name].abs().sort_values(ascending=False).head(3)
    print(f"\n{pc_name} ({explained_variance[i]*100:.2f}% variance):")
    for feat, loading in top_features.items():
        print(f"  {feat}: {pca_loadings.loc[feat, pc_name]:.4f}")

# ============== 3. FEATURE IMPORTANCE (XGBoost) ==============
print("\n" + "=" * 70)
print("3. FEATURE IMPORTANCE (XGBoost)")
print("=" * 70)

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)

model = XGBClassifier(n_estimators=100, max_depth=6, learning_rate=0.1, random_state=42, use_label_encoder=False, eval_metric='mlogloss')
model.fit(X_train, y_train)

# Feature importance
importance = pd.DataFrame({
    'Feature': features,
    'Importance': model.feature_importances_
}).sort_values('Importance', ascending=False)

print("\nFeature Importance (XGBoost):")
for _, row in importance.iterrows():
    print(f"  {row['Feature']}: {row['Importance']:.4f}")

# ============== 4. RECOMMENDATIONS ==============
print("\n" + "=" * 70)
print("4. FEATURE RECOMMENDATIONS")
print("=" * 70)

# Features with low importance (< 0.05)
low_importance = importance[importance['Importance'] < 0.05]['Feature'].tolist()

# Features with high correlation (> 0.8)
high_corr_features = [pair[1] for pair in high_corr_pairs]

# Features to potentially remove
features_to_remove = set(low_importance + high_corr_features)
features_to_keep = [f for f in features if f not in features_to_remove]

print(f"\nFeatures with LOW importance (< 0.05): {low_importance}")
print(f"HIGHLY correlated features: {high_corr_features}")
print(f"\nFeatures RECOMMENDED TO REMOVE: {list(features_to_remove)}")
print(f"Features RECOMMENDED TO KEEP: {features_to_keep}")

# ============== 5. MODEL COMPARISON (Original vs Reduced) ==============
print("\n" + "=" * 70)
print("5. MODEL PERFORMANCE: ORIGINAL vs REDUCED FEATURES")
print("=" * 70)

# Original model (all features)
y_pred_original = model.predict(X_test)
accuracy_original = accuracy_score(y_test, y_pred_original)

# Reduced model (remove low importance features)
X_reduced = df_clean[features_to_keep]
X_train_red, X_test_red, y_train_red, y_test_red = train_test_split(
    X_reduced, y, test_size=0.2, random_state=42, stratify=y
)

model_red = XGBClassifier(n_estimators=100, max_depth=6, learning_rate=0.1, random_state=42, use_label_encoder=False, eval_metric='mlogloss')
model_red.fit(X_train_red, y_train_red)

y_pred_reduced = model_red.predict(X_test_red)
accuracy_reduced = accuracy_score(y_test_red, y_pred_reduced)

print(f"\nOriginal Model (12 features): {accuracy_original*100:.2f}%")
print(f"Reduced Model ({len(features_to_keep)} features): {accuracy_reduced*100:.2f}%")
print(f"Accuracy difference: {(accuracy_original - accuracy_reduced)*100:.2f}%")

# ============== 6. CREATE VISUALIZATIONS ==============

# Plot 1: Correlation Heatmap
plt.figure(figsize=(12, 10))
sns.heatmap(corr_matrix, annot=True, cmap='RdBu_r', center=0, fmt='.2f')
plt.title('Feature Correlation Matrix')
plt.tight_layout()
plt.savefig('correlation_heatmap.png', dpi=150)
plt.close()
print("\nSaved: correlation_heatmap.png")

# Plot 2: PCA Explained Variance
plt.figure(figsize=(10, 6))
plt.bar(range(1, len(explained_variance)+1), explained_variance*100, alpha=0.7, label='Individual')
plt.plot(range(1, len(cumulative_variance)+1), cumulative_variance*100, 'ro-', label='Cumulative')
plt.axhline(y=95, color='g', linestyle='--', label='95% threshold')
plt.xlabel('Principal Component')
plt.ylabel('Explained Variance (%)')
plt.title('PCA - Explained Variance')
plt.legend()
plt.grid(True, alpha=0.3)
plt.tight_layout()
plt.savefig('pca_variance.png', dpi=150)
plt.close()
print("Saved: pca_variance.png")

# Plot 3: Feature Importance
plt.figure(figsize=(10, 6))
colors = ['green' if imp >= 0.05 else 'red' for imp in importance['Importance']]
plt.barh(importance['Feature'], importance['Importance'], color=colors)
plt.xlabel('Importance')
plt.title('Feature Importance (XGBoost)')
plt.axvline(x=0.05, color='black', linestyle='--', label='Threshold (0.05)')
plt.tight_layout()
plt.savefig('feature_importance.png', dpi=150)
plt.close()
print("Saved: feature_importance.png")

# ============== 7. CREATE WORD DOCUMENT ==============
print("\n" + "=" * 70)
print("CREATING WORD DOCUMENT")
print("=" * 70)

doc = Document()

# Title
title = doc.add_heading('PCA Analysis Report - Sleep Quality Dataset', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Dataset Overview
doc.add_heading('1. Dataset Overview', level=1)
doc.add_paragraph(f'Original samples: {len(df)}')
doc.add_paragraph(f'Cleaned samples: {len(df_clean)}')
doc.add_paragraph(f'Features analyzed: {len(features)}')
doc.add_paragraph(f'Features: {", ".join(features)}')

# Correlation Analysis
doc.add_heading('2. Correlation Analysis', level=1)
doc.add_paragraph('Correlation with Sleep Quality (target):')
for feat, corr in corr_with_target.items():
    doc.add_paragraph(f'  • {feat}: {corr:.4f}')

if high_corr_pairs:
    doc.add_paragraph('\nHighly Correlated Feature Pairs (|r| > 0.8):')
    for pair in high_corr_pairs:
        doc.add_paragraph(f"  • {pair[0]} <-> {pair[1]}: {pair[2]:.4f}")
else:
    doc.add_paragraph('\nNo highly correlated feature pairs found.')

doc.add_picture('correlation_heatmap.png', width=Inches(5.5))

# PCA Results
doc.add_heading('3. PCA Results', level=1)
doc.add_paragraph('Explained Variance by Component:')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Variance %'
hdr_cells[2].text = 'Cumulative %'
hdr_cells[3].text = 'Key Features'

for i in range(min(6, len(features))):
    row_cells = table.add_row().cells
    row_cells[0].text = f'PC{i+1}'
    row_cells[1].text = f'{explained_variance[i]*100:.2f}%'
    row_cells[2].text = f'{cumulative_variance[i]*100:.2f}%'

    # Top 2 features
    top_feats = pca_loadings[f'PC{i+1}'].abs().nlargest(2).index.tolist()
    row_cells[3].text = ', '.join(top_feats)

doc.add_paragraph(f'\nComponents for 90% variance: {n_components_90}')
doc.add_paragraph(f'Components for 95% variance: {n_components_95}')

doc.add_picture('pca_variance.png', width=Inches(5.5))

# Feature Importance
doc.add_heading('4. Feature Importance (XGBoost)', level=1)

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Importance'

for _, row in importance.iterrows():
    row_cells = table2.add_row().cells
    row_cells[0].text = row['Feature']
    row_cells[1].text = f'{row["Importance"]:.4f}'

doc.add_picture('feature_importance.png', width=Inches(5.5))

# Recommendations
doc.add_heading('5. Feature Recommendations', level=1)

doc.add_paragraph('Features to REMOVE (low importance or high correlation):')
for feat in features_to_remove:
    doc.add_paragraph(f"  • {feat}")

doc.add_paragraph('\nFeatures to KEEP:')
for feat in features_to_keep:
    doc.add_paragraph(f"  • {feat}")

# Model Comparison
doc.add_heading('6. Model Performance Comparison', level=1)
doc.add_paragraph(f'Original Model (all {len(features)} features): {accuracy_original*100:.2f}%')
doc.add_paragraph(f'Reduced Model ({len(features_to_keep)} features): {accuracy_reduced*100:.2f}%')
doc.add_paragraph(f'Accuracy difference: {(accuracy_original - accuracy_reduced)*100:.2f}%')

# Conclusion
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    f'Based on the PCA and feature importance analysis, the dataset can be reduced from '
    f'{len(features)} to {len(features_to_keep)} features with minimal accuracy loss. '
    f'The most important features for sleep quality prediction are: '
    f'{", ".join(importance.head(5)["Feature"].tolist())}.'
)

doc.save('PCA_analysis.docx')
print("\nWord document saved: PCA_analysis.docx")

print("\n" + "=" * 70)
print("PCA ANALYSIS COMPLETE")
print("=" * 70)